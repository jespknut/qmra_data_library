import openai
import os

from openai import OpenAI

client = OpenAI(api_key="")  # Paste your real key here

def generate_ai_summary(prompt_text):
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a professional water reuse consultant writing interpretive risk summaries."},
            {"role": "user", "content": prompt_text}
        ],
        temperature=0.5
    )
    return response.choices[0].message.content


from docx import Document
from docx.shared import Inches
import tempfile

def create_docx_report(summary_text, ai_summary, fig1, fig2):
    doc = Document()
    doc.add_heading("QMRA Risk Assessment Report", 0)

    doc.add_heading("Scenario Summary", level=1)
    for line in summary_text.strip().split("\n"):
        doc.add_paragraph(line.strip())

    if ai_summary:
        doc.add_heading("AI-Generated Interpretation", level=1)
        for line in ai_summary.strip().split("\n"):
            doc.add_paragraph(line.strip())

    doc.add_heading("Figures", level=1)

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp1, \
         tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp2:

        fig1.savefig(tmp1.name, dpi=300, bbox_inches='tight')
        fig2.savefig(tmp2.name, dpi=300, bbox_inches='tight')

        doc.add_picture(tmp1.name, width=Inches(5.5))
        doc.add_picture(tmp2.name, width=Inches(5.5))

    # Save DOCX to a temporary file
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_docx.name)
    tmp_docx.seek(0)
    return tmp_docx



import streamlit as st

# Load data from your QMRA script
from treatment_Scenarios_QMRA_greywater_1_2 import (
    SOURCE_WATER, PATHOGENS, ACTIVITIES, TREATMENT_EFFECTIVENESS_DIST
)

# Title
st.title("QMRA for water reuse - Scenario Builder")

# Sidebar inputs
st.sidebar.header("1. Water Source")
water_source = st.sidebar.selectbox("Select water source", list(SOURCE_WATER.keys()))

st.sidebar.header("2. Reuse Activities")
selected_activities = st.sidebar.multiselect(
    "Select one or more reuse purposes",
    [a['Name'] for a in ACTIVITIES],
    default=["Shower"]
)

st.sidebar.header("3. Pathogens of Concern")
selected_pathogens = st.sidebar.multiselect(
    "Select pathogens to simulate",
    [p['Name'] for p in PATHOGENS],
    default=["Rotavirus", "Salmonella spp."]
)

st.sidebar.header("4. Treatment Train")
treatment_options = list(TREATMENT_EFFECTIVENESS_DIST.keys())
selected_treatments = st.sidebar.multiselect(
    "Select treatment steps (in order)",
    treatment_options,
    default=["Sand filter", "UV"]
)

# Add treatment status selection
treatment_status = {}
for step in selected_treatments:
    status = st.sidebar.selectbox(
        f"Status of '{step}'",
        options=["normal", "degraded", "fail"],
        key=f"status_{step}"
    )
    treatment_status[step] = status

st.sidebar.header("5. Exposure Population")
population_size = st.sidebar.number_input("Number of exposed users", min_value=1, value=10, step=1)

# Placeholder: Modifiers (future)
modifiers = st.sidebar.multiselect(
    "Optional risk modifiers (future)",
    ["Children only", "Immunocompromised", "Conservative (50% ↑ exposure)", "Behavioral reduction (50% ↓)"]
)

# --- Show Summary ---
st.subheader("🧾 Scenario Summary")
st.markdown(f"""
**Water Source:** {water_source}  
**Activities:** {', '.join(selected_activities)}  
**Pathogens:** {', '.join(selected_pathogens)}  
**Treatment Train:**  
""" + '\n'.join([f"- {t} ({treatment_status[t]})" for t in selected_treatments]) + f"""
\n**Users:** {population_size}
""")

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import altair as alt
from treatment_Scenarios_QMRA_greywater_1_2 import simulate_mc, PATHOGENS, ACTIVITIES

# --- Prepare treatment sequence ---
treatment_steps = [(name, None, treatment_status[name]) for name in selected_treatments]

# --- Run simulations ---
st.subheader("📊 Simulation Results")

# Add button to trigger simulation
if st.button("Run QMRA Simulation"):

    all_results = []

    with st.spinner("Running simulations..."):
        for activity_name in selected_activities:
            activity = next(a for a in ACTIVITIES if a["Name"] == activity_name)

            for pathogen_name in selected_pathogens:
                pathogen = next(p for p in PATHOGENS if p["Name"] == pathogen_name)

                df = simulate_mc(
                    activity=activity,
                    pathogen=pathogen,
                    treatment_steps=treatment_steps,
                    scenario_name="Custom scenario",
                    source=water_source,
                    population=population_size
                )
                all_results.append(df)

    if all_results:
        combined_df = pd.concat(all_results, ignore_index=True)
       # Collect pathogen DALYs into separate series
        all_daly = []
        for pathogen in selected_pathogens:
            daly_values = combined_df.loc[combined_df["Pathogen"] == pathogen, "DALY"].reset_index(drop=True)
            all_daly.append(daly_values)

        # Build single DataFrame
        total_daly_df = pd.concat(all_daly, axis=1)
        total_daly_df.columns = selected_pathogens

        # Sum across pathogens per row
        total_daly = total_daly_df.sum(axis=1)
        
               
        # Compute total Expected cases, Annual risk, and Infection probability
        all_expected_cases = []
        all_annual_risk = []
        all_pinf = []

        for pathogen in selected_pathogens:
            # Filter rows for this pathogen
            df_p = combined_df[combined_df["Pathogen"] == pathogen]
            all_expected_cases.append(df_p["Expected cases"].reset_index(drop=True))
            all_annual_risk.append(df_p["Annual risk"].reset_index(drop=True))
            all_pinf.append(df_p["P_inf"].reset_index(drop=True))

        # Build DataFrames
        expected_cases_df = pd.concat(all_expected_cases, axis=1)
        annual_risk_df = pd.concat(all_annual_risk, axis=1)
        pinf_df = pd.concat(all_pinf, axis=1)

        # Add column names
        expected_cases_df.columns = selected_pathogens
        annual_risk_df.columns = selected_pathogens
        pinf_df.columns = selected_pathogens

        # Sum across pathogens for totals
        total_expected_cases = expected_cases_df.sum(axis=1)
        total_annual_risk = annual_risk_df.sum(axis=1)
        total_pinf = pinf_df.sum(axis=1)

        # For pathogen-wise plots
        daly_by_pathogen = total_daly_df

        
        # --- Auto-Generated Report Summary ---
        st.markdown("### 📝 Auto-generated Summary")

        p50 = np.percentile(total_daly, 50)
        p95 = np.percentile(total_daly, 95)
        
        p50_cases = np.percentile(total_expected_cases, 50)
        p95_cases = np.percentile(total_expected_cases, 95)

        p50_pinf = np.percentile(total_pinf, 50)
        p95_pinf = np.percentile(total_pinf, 95)

        p50_risk = np.percentile(total_annual_risk, 50)
        p95_risk = np.percentile(total_annual_risk, 95)
        
        exceeds = p95 > 1e-6

        summary_text = f"""
        This simulation includes {population_size} exposed users across {len(selected_activities)} reuse activity(ies) using {water_source.lower()} as source water.

        The treatment train consisted of:
        {', '.join([f"{step} ({treatment_status[step]})" for step in selected_treatments])}.

        **Median total DALY per person-year:** {p50:.2e}
        95th percentile DALY: {p95:.2e}

        **Median total expected cases per person-year:** {p50_cases:.2e}
        95th percentile expected cases: {p95_cases:.2e}

        **Median total infection probability:** {p50_pinf:.2e}
        95th percentile infection probability: {p95_pinf:.2e}

        **Median total annual risk of infection:** {p50_risk:.2e}
        95th percentile annual risk: {p95_risk:.2e}

        This {'exceeds' if exceeds else 'meets'} the WHO benchmark of **1×10⁻⁶ DALY/person-year**.
        """

        st.text_area("Summary for reporting", summary_text.strip(), height=180)


        # Show DALY stats
        st.markdown("### Mean DALY per pathogen and activity")
        mean_daly = (
            combined_df.groupby(["Pathogen", "Activity"])["DALY"]
            .mean()
            .reset_index()
            .sort_values(by="DALY", ascending=False)
        )
        st.dataframe(mean_daly.style.format({"DALY": "{:.2e}"}))

        # --- DALY histogram with log-scale x-axis and better binning ---
        st.markdown("### Total DALY Distribution")

        # Drop zero or negative values to avoid log issues
        dalys = total_daly[total_daly > 0]

        if dalys.empty:
            st.warning("No DALY values above zero to plot.")
        else:
            # Percentiles
            p5 = np.percentile(dalys, 5)
            p50 = np.percentile(dalys, 50)
            p95 = np.percentile(dalys, 95)

            # Store in session or for reuse later
            st.session_state['p50'] = p50
            st.session_state['p95'] = p95

            # Binning
            min_val = max(dalys.min(), 1e-9)
            max_val = dalys.max() * 1.1
            bins = np.logspace(np.log10(min_val), np.log10(max_val), 50)

            fig, ax = plt.subplots(figsize=(8, 4))
            ax.hist(dalys, bins=bins, color='orchid', edgecolor='black', alpha=0.7)

            # WHO threshold region
            ax.axvspan(0, 1e-6, color='red', alpha=0.2, label='WHO benchmark (1e-6)')

            # Annotate key percentiles
            ax.axvline(p50, color='blue', linestyle='--', label=f'Median: {p50:.1e}')
            ax.axvline(p95, color='green', linestyle='--', label=f'95th percentile: {p95:.1e}')

            ax.set_xscale('log')
            ax.set_title("Total DALY per person-year")
            ax.set_xlabel("DALY (log scale)")
            ax.set_ylabel("Frequency")
            ax.legend()
            st.pyplot(fig)

            st.caption(f"{(total_daly < 1e-9).sum()} results were < 1e-9 and may not show clearly on a log scale.")

        
        # --- Log-scale Pathogen Contribution Bar Chart ---
        st.markdown("### Pathogen Contributions to Total DALY")

        mean_contrib = (
            combined_df.groupby("Pathogen")["DALY"]
            .mean()
            .sort_values(ascending=False)
        )

        fig2, ax2 = plt.subplots(figsize=(10, 5))
        mean_contrib.plot(kind='bar', color='teal', edgecolor='black', ax=ax2)
        ax2.set_yscale('log')
        ax2.set_ylabel("Mean DALY (log scale)")
        ax2.set_title("Mean Pathogen DALY Contribution (log scale)")
        ax2.set_xlabel("Pathogen")
        ax2.set_xticklabels(mean_contrib.index, rotation=45, ha='right')
        st.pyplot(fig2)
        
        import seaborn as sns

        if "total_daly_by_pathogen" in st.session_state:
            st.markdown("### 🔍 Pathogen-wise DALY Distributions")

            fig3, ax3 = plt.subplots(figsize=(10, 5))
            for pathogen in st.session_state["total_daly_by_pathogen"].columns:
                sns.kdeplot(
                    st.session_state["total_daly_by_pathogen"][pathogen],
                    ax=ax3,
                    label=pathogen,
                    bw_adjust=0.5,
                    clip=(1e-10, 1e-3)
                )

            ax3.set_xscale("log")
            ax3.set_xlabel("DALY (log scale)")
            ax3.set_title("Pathogen-wise DALY Distributions")
            ax3.legend()
            st.pyplot(fig3)
        
        
        st.markdown("## 🔍 Additional Risk Distributions")

        # Plot total expected cases
        if total_expected_cases.max() > 0:
            fig_cases, ax_cases = plt.subplots(figsize=(8, 4))
            ax_cases.hist(total_expected_cases, bins=40, color='skyblue', edgecolor='black', alpha=0.7)
            ax_cases.set_title("Total Expected Cases (all pathogens)")
            ax_cases.set_xlabel("Expected Cases per Person-Year")
            ax_cases.set_ylabel("Frequency")
            st.pyplot(fig_cases)

        # Plot total infection probability
        if total_pinf.max() > 0:
            fig_pinf, ax_pinf = plt.subplots(figsize=(8, 4))
            ax_pinf.hist(total_pinf, bins=40, color='lightcoral', edgecolor='black', alpha=0.7)
            ax_pinf.set_title("Total Infection Probability (all pathogens)")
            ax_pinf.set_xlabel("Probability")
            ax_pinf.set_ylabel("Frequency")
            st.pyplot(fig_pinf)

        # Plot total annual risk
        if total_annual_risk.max() > 0:
            fig_risk, ax_risk = plt.subplots(figsize=(8, 4))
            ax_risk.hist(total_annual_risk, bins=40, color='lightgreen', edgecolor='black', alpha=0.7)
            ax_risk.set_title("Total Annual Risk (all pathogens)")
            ax_risk.set_xlabel("Annual Probability of Infection")
            ax_risk.set_ylabel("Frequency")
            st.pyplot(fig_risk)
        
        # --- AI summary prompt generation ---
        st.markdown("### 🤖 AI-Powered Report Summary")

        top_contributors = mean_contrib.head(3).to_string()

        prompt = f"""
        Write a summary for a quantitative microbial risk assessment scenario.

        Water source: {water_source}
        Reuse activities: {', '.join(selected_activities)}
        Pathogens simulated: {', '.join(selected_pathogens)}
        Treatment train: {', '.join([f"{s} ({treatment_status[s]})" for s in selected_treatments])}
        Number of exposed users: {population_size}

        Median total DALY: {p50:.2e}
        95th percentile DALY: {p95:.2e}

        Median total expected cases per person-year: {p50_cases:.2e}
        95th percentile expected cases: {p95_cases:.2e}

        Median total infection probability: {p50_pinf:.2e}
        95th percentile infection probability: {p95_pinf:.2e}

        Median total annual risk of infection: {p50_risk:.2e}
        95th percentile annual risk: {p95_risk:.2e}

        WHO benchmark: 1e-6 DALY/person-year

        Top 3 contributing pathogens:
        {top_contributors}

        Interpret this in a professional tone and offer practical recommendations.
        """

        try:
            with st.spinner("Generating AI summary..."):
                ai_summary = generate_ai_summary(prompt)
                st.success("Summary generated.")
                st.text_area("AI-generated report summary", ai_summary, height=300)
        except Exception as e:
            st.error(f"AI summary generation failed: {e}")



        # Download button
        csv = combined_df.to_csv(index=False)
        st.download_button("Download full results (CSV)", csv, "qmra_results.csv")
        
        if 'summary_text' in locals() and 'ai_summary' in locals():
            docx_file = create_docx_report(summary_text, ai_summary, fig, fig2)

            with open(docx_file.name, "rb") as f:
                st.download_button(
                    label="📄 Download DOCX Report",
                    data=f.read(),
                    file_name="qmra_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


    else:
        st.warning("No results generated.")


